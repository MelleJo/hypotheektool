import streamlit as st
import os
from PyPDF2 import PdfReader
from docx import Document
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain.chains import AnalyzeDocumentChain
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate

def extract_text_from_pdf(pdf_path):
    reader = PdfReader(pdf_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text])
    return text

def generate_report(user_input, motivational_texts, template_text):
    with st.spinner('Verwerken...'):
        # Prepare the AI prompt
        prompt = f"""
        Hier is de input van de gebruiker over een hypotheekadviesproces: "{user_input}". Gebruik deze informatie samen met de onderstaande motivatieteksten en sjabloon om een volledig gestructureerd hypotheekadviesrapport te genereren. Het doel is om een conceptdocument te creëren dat later kan worden aangepast met specifieke namen en details door de gebruiker.

        === Sjabloon Inhoud ===
        {template_text}
        
        === Motivatieteksten ===
        {motivational_texts}

        Genereer op basis van deze gegevens een coherent adviesrapport in tekstvorm.
        """

        # Set up the LangChain LLM chain
        llm = ChatOpenAI(api_key=st.secrets["OPENAI_API_KEY"], model="gpt-4-turbo-2024-04-09", temperature=0, streaming=True)
        chain = ChatPromptTemplate(prompt) | llm | StrOutputParser()

        # Execute chain
        output = chain.stream()

        # Save output to a .docx file
        output_doc = Document()
        output_doc.add_paragraph(output)
        output_path = os.path.join('output', 'Hypotheek_Rapport.docx')
        output_doc.save(output_path)
        return output_path


def main():
    st.title("Hypotheektool - Testversie 0.0.1")
    user_input = st.text_input("Geef informatie over het hypotheek adviestraject. Voeg hier geen persoonlijk identificeerbare informatie toe; deze details kun je later toevoegen.")
    motivational_texts = extract_text_from_docx("source documents/Teksten voor hypotheekadviesrapport in Fastlane (in concept, MvdB).docx")
    template_text = extract_text_from_pdf("source documents/Hyp. rapport zonder motivatieteksten.pdf")
    
    if st.button("Rapport Genereren"):
        report_path = generate_report(user_input, motivational_texts, template_text)
        if os.path.exists(report_path):
            with open(report_path, "rb") as file:
                st.success('Rapport succesvol gegenereerd.')
                st.download_button(
                    label="Download Rapport",
                    data=file,
                    file_name="Hypotheek_Rapport.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.error("Er is een fout opgetreden bij het genereren van het rapport.")

if __name__ == "__main__":
    st.set_page_config(page_title="Hypotheek Adviesrapport Tool")
    main()
